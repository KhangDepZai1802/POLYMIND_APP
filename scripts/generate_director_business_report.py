from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "POLYMIND_Bao_cao_quy_trinh_nghiep_vu_cho_Giam_doc_2026-07-11.docx"
LOGO = ROOT / "POLYMIND.png"

BLUE = "086BC1"
DARK_BLUE = "17365D"
ORANGE = "F28C28"
LIGHT_BLUE = "EAF4FC"
LIGHT_ORANGE = "FFF2E5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "1F7A4D"
RED = "B42318"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.size = Pt(8)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mở file bằng Microsoft Word và chọn Update Field để cập nhật mục lục."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, placeholder, end])


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, color: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False if widths else True
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, DARK_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        if widths:
            cell.width = Cm(widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(9)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = Cm(widths[i])
            if ridx % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            run.font.size = Pt(8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in (("Title", 28, DARK_BLUE), ("Heading 1", 18, DARK_BLUE),
                              ("Heading 2", 13.5, BLUE), ("Heading 3", 11.5, ORANGE)):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.text = "POLYMIND OLMS • Báo cáo quy trình & nghiệp vụ chức năng • 11/07/2026"
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_page_number(section.footer.add_paragraph())

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    return doc


def build_report() -> Document:
    doc = setup_document()

    # Cover
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(3.7))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("BÁO CÁO TỔNG QUAN\nQUY TRÌNH & NGHIỆP VỤ CHỨC NĂNG")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("HỆ THỐNG POLYMIND OLMS")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(ORANGE)
    doc.add_paragraph()
    add_callout(doc, "Dành cho Ban Giám đốc", "Báo cáo điều hành về chuỗi nghiệp vụ từ Lead đến hoàn tất nghĩa vụ sau xuất cảnh, phạm vi chức năng hiện có, cơ chế kiểm soát, chỉ số quản trị và các quyết định cần ưu tiên.", LIGHT_BLUE)
    doc.add_paragraph()
    cover_meta = add_table(doc, ["Thông tin", "Nội dung"], [
        ["Ngày báo cáo", "11/07/2026"],
        ["Phạm vi rà soát", "Toàn bộ web hiện tại: mã nguồn, màn hình, quy tắc nghiệp vụ, dữ liệu miền, tài liệu QA và kiểm thử"],
        ["Trạng thái xác minh", "88/88 kiểm thử tự động đạt trên mã nguồn hiện tại"],
        ["Mục đích", "Cung cấp bức tranh tổng thể để điều hành, phê duyệt và ưu tiên hoàn thiện trước khi vận hành thật"],
    ], [4.3, 12.6])
    doc.add_paragraph()
    p = doc.add_paragraph("Tài liệu nội bộ — POLYMIND / Vietgroup Edu")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor.from_string(MID_GRAY)
    doc.add_page_break()

    add_heading(doc, "Mục lục", 1)
    add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "1. Tóm tắt điều hành", 1)
    doc.add_paragraph(
        "POLYMIND OLMS hiện là nền tảng quản lý tập trung cho hoạt động tuyển dụng và đưa người lao động/học viên đi làm việc, học tập hoặc tham gia việc làm trong và ngoài nước. Hệ thống bao phủ chuỗi giá trị từ tiếp nhận Lead, tư vấn, tạo hồ sơ ứng viên, gắn đơn hàng, đào tạo, thu phí, hỗ trợ vay, visa–vé máy bay, xuất cảnh, chăm sóc tại nước ngoài đến hoàn tất nghĩa vụ."
    )
    add_callout(doc, "Kết luận chính", "Phạm vi chức năng cốt lõi đã hình thành tương đối đầy đủ và có liên kết dữ liệu xuyên suốt. Giá trị lớn nhất với Ban Giám đốc là nhìn được phễu tuyển dụng, doanh thu–chi phí–công nợ, tiến độ ứng viên, hiệu suất kênh/đại lý và các điểm nghẽn vận hành trên cùng một hệ thống.", LIGHT_BLUE)
    add_heading(doc, "1.1 Năm năng lực nổi bật", 2)
    add_bullets(doc, [
        "Một hồ sơ xuyên suốt: Lead chuyển thành ứng viên và tiếp tục đi qua quy trình 20 bước, hạn chế nhập lại dữ liệu.",
        "Kiểm soát theo vai trò và phạm vi dữ liệu: nhân sự nội bộ, đại lý, CTV, phụ huynh và học viên nhìn các phần khác nhau.",
        "Liên kết tiến độ với tài chính và hoa hồng: thu tiền theo giai đoạn, hoa hồng sinh theo mốc, công nợ và khoản vay ảnh hưởng điều kiện hoàn tất.",
        "Dashboard và báo cáo quản trị: theo dõi phễu, tỷ lệ chuyển đổi, trúng tuyển, đậu visa, xuất cảnh, doanh thu, chi phí, lợi nhuận gộp và công nợ.",
        "Hệ sinh thái hỗ trợ vận hành: thông báo, nhắn tin, lưu hồ sơ, nhật ký audit, xuất Excel/PDF/CSV và trợ lý AI."
    ])
    add_heading(doc, "1.2 Đánh giá mức độ sẵn sàng", 2)
    add_table(doc, ["Nhóm", "Đánh giá hiện tại", "Ý nghĩa điều hành"], [
        ["Nền tảng cốt lõi M01–M12", "Phần lớn đã QA ở mức mã nguồn; các lỗi quan trọng đã được sửa. M11 vừa sửa và còn chờ xác minh độc lập/runtime.", "Có thể tiếp tục UAT theo vai trò và dữ liệu test có kiểm soát."],
        ["Thông báo M13", "Đang còn quyết định nghiệp vụ về người nhận và mức thông tin dành cho CTV.", "Cần Giám đốc chốt chính sách trước khi đóng module."],
        ["M14–M20", "Đã có chức năng ở web nhưng QA tổng thể còn đang chờ/thực hiện từng phần.", "Chưa nên xem là hoàn tất kiểm định production."],
        ["Kiểm thử hiện tại", "88/88 test tự động đạt; nhiều xác minh mới ở code-level, E2E/DB/UI chưa phủ đầy đủ.", "Kết quả tốt nhưng không thay thế UAT và kiểm thử vận hành thật."],
    ], [3.2, 7.2, 6.5])
    add_callout(doc, "Khuyến nghị điều hành 30 ngày", "Chốt các quy tắc còn mở; hoàn tất UAT theo 12 vai trò; thử xuyên suốt 5 hồ sơ mẫu từ Lead đến xuất cảnh; áp migration trên DB test; chạy kiểm thử tải, bảo mật, backup/restore và nghiệm thu báo cáo trước khi public production.", LIGHT_ORANGE)

    add_heading(doc, "2. Phạm vi và phương pháp rà soát", 1)
    doc.add_paragraph("Báo cáo được tổng hợp từ trạng thái web và tài liệu tại ngày 11/07/2026, không chỉ dựa trên tài liệu đặc tả cũ.")
    add_table(doc, ["Nguồn kiểm tra", "Quy mô/điểm kiểm"], [
        ["Giao diện Blazor", "49 file trang/component nghiệp vụ trong thư mục Pages; kiểm route, menu, nút thao tác, điều kiện quyền và luồng điều hướng."],
        ["Mô hình dữ liệu", "27 entity miền; kiểm liên kết Lead–Candidate–Job–Workflow–Finance–Loan–Visa–Commission–Notification."],
        ["Quy tắc và dịch vụ", "Workflow 20 bước, chính sách quyền, tài chính, hoa hồng, vay/nợ, thông báo, nhắn tin, AI và lưu trữ hồ sơ."],
        ["Tài liệu QA", "Bảng QA 20 module và hồ sơ phân tích/business flow/test/fix/verification của M01–M14."],
        ["Kiểm thử", "Chạy trực tiếp dotnet test: 88 đạt, 0 lỗi, 0 bỏ qua."],
    ], [4.2, 12.7])
    doc.add_paragraph("Lưu ý: tài liệu cũ mô tả 17 bước; báo cáo này dùng quy trình 20 bước trong mã nguồn hiện tại làm nguồn chuẩn.")

    add_heading(doc, "3. Bức tranh quy trình nghiệp vụ tổng thể", 1)
    add_callout(doc, "Chuỗi giá trị end-to-end", "Nguồn Lead → CRM & tư vấn → Chuyển thành ứng viên → Chọn đơn hàng → Hồ sơ/khám sức khỏe/thi tuyển → Đào tạo & hợp đồng → COE/Visa → Thu phí/vay/công nợ → Vé máy bay/xuất cảnh → Hỗ trợ tại nước ngoài → Hoàn tất nghĩa vụ.")
    add_heading(doc, "3.1 Quy trình cấp cao", 2)
    add_table(doc, ["Giai đoạn", "Đầu vào", "Xử lý chính", "Đầu ra/kiểm soát"], [
        ["1. Thu hút & tiếp nhận", "Lead từ Facebook, TikTok, Google, Website, Zalo, hotline, đại lý, giới thiệu, sự kiện", "Tạo Lead, gán tư vấn viên, ghi hoạt động, hẹn tư vấn, nhắc chăm sóc quá hạn", "Lead đủ thông tin và có người chịu trách nhiệm"],
        ["2. Tư vấn & chuyển đổi", "Lead quan tâm/đăng ký", "Cập nhật 10 trạng thái; chuyển Lead thành ứng viên khi đủ điều kiện", "Hồ sơ ứng viên duy nhất, giữ nguồn và người giới thiệu"],
        ["3. Hồ sơ & đơn hàng", "Ứng viên + nhu cầu/thị trường", "Bổ sung thông tin, tài liệu, gắn Job; đổi Job đặc biệt phải super admin xác nhận mật khẩu", "Ứng viên có đơn hàng và workflow active"],
        ["4. Tuyển chọn & đào tạo", "Hồ sơ, sức khỏe, đơn hàng", "Khám sức khỏe, thi/phỏng vấn, chọn lại Job nếu trượt, học tiếng/định hướng/nghề, đánh giá định kỳ", "Ứng viên trúng tuyển và sẵn sàng ký hợp đồng"],
        ["5. Pháp lý & tài chính", "Hợp đồng, COE, hồ sơ visa, lịch phí", "Thu tiền 4 đợt; quản lý chi; hỗ trợ vay/nợ công ty; sinh phiếu; theo dõi công nợ", "Đủ điều kiện pháp lý và tài chính trước xuất cảnh"],
        ["6. Xuất cảnh & sau xuất cảnh", "Visa đạt, thanh toán/cam kết nợ, vé bay", "Theo dõi chuyến bay, xuất cảnh, nhật ký hỗ trợ tại nước ngoài", "Hoàn thành khi hết nghĩa vụ; nợ công ty phải thu đủ"],
    ], [2.8, 3.8, 6.0, 4.4])

    add_heading(doc, "3.2 Quy trình 20 bước của ứng viên", 2)
    workflow_rows = [
        ["1", "Lead mới", "Tuyển dụng/CRM"], ["2", "Đã liên hệ / liên hệ lại", "Tuyển dụng/CRM"],
        ["3", "Đã tư vấn", "Tư vấn"], ["4", "Đăng ký", "Tư vấn"], ["5", "Đặt cọc", "Tài chính"],
        ["6", "Hoàn thiện hồ sơ", "Hồ sơ"], ["7", "Khám sức khỏe", "Hồ sơ"],
        ["7.5", "Chọn lại đơn hàng khác nếu rớt bước 8", "Tuyển dụng"],
        ["8", "Thi tuyển / phỏng vấn / xét duyệt hồ sơ", "Tuyển dụng + Hồ sơ"], ["9", "Trúng tuyển", "Tuyển dụng"],
        ["10", "Học tiếng / định hướng / nghề (nếu cần)", "Đào tạo/Hồ sơ"], ["11", "Ký hợp đồng", "Hồ sơ"],
        ["12", "Xin COE / tư cách lưu trú", "Hồ sơ/Visa"], ["13", "Nộp hồ sơ Visa", "Visa"],
        ["14", "Đậu Visa", "Visa"], ["15", "Thanh toán hoàn tất / cam kết trả nợ", "Kế toán"],
        ["16", "Đặt vé máy bay", "Visa/Xuất cảnh"], ["17", "Xuất cảnh", "Visa/Xuất cảnh"],
        ["18", "Đến nơi làm việc", "Vận hành"], ["19", "Tương tác giai đoạn xứ người", "Tư vấn/Hồ sơ"],
        ["20", "Hoàn thành quy trình", "Kiểm soát nghĩa vụ"],
    ]
    add_table(doc, ["Bước", "Trạng thái nghiệp vụ", "Bộ phận chính"], workflow_rows, [1.7, 10.0, 5.2])
    add_bullets(doc, [
        "Chuyển bước phải đúng nhóm phụ trách; hệ thống lưu lịch sử bước, trạng thái, người thực hiện và thời gian.",
        "Nếu rớt bước 8, ứng viên vào bước phụ 7.5 và bắt buộc chọn đơn hàng mới còn hạn trước khi thi lại.",
        "Bước 19 cho phép ghi nhiều nhật ký hỗ trợ trong nhiều năm, không buộc đóng hồ sơ ngay.",
        "Bước 20 bị chặn nếu còn khoản nợ công ty chưa thu đủ; không có cơ chế miễn nợ.",
        "Super admin đổi đơn hàng có thể reset toàn bộ workflow về bước 1, nhưng không tự động hoàn khoản thu/hoa hồng đã phát sinh theo quyết định hiện hành."
    ])

    add_heading(doc, "4. Danh mục chức năng hiện có", 1)
    module_rows = [
        ["1", "Đăng nhập & phiên", "Đăng nhập web/API, khóa tạm, vô hiệu hóa tài khoản, đăng xuất, kiểm tra lại phiên."],
        ["2", "Vai trò & phân quyền", "RBAC 20 tài nguyên × 5 hành động; phạm vi dữ liệu theo đại lý/CTV/cá nhân."],
        ["3", "Quản trị tài khoản", "Tạo, sửa, khóa/mở, đổi role, reset/tự đổi mật khẩu, tài khoản đối tác/phụ huynh/học viên."],
        ["4", "Lead CRM", "Nguồn Lead, 10 trạng thái, hoạt động chăm sóc, phân công, hẹn, convert/revert, nhắc quá hạn."],
        ["5", "Ứng viên & hồ sơ", "Thông tin cá nhân, giấy tờ, tài khoản cổng, người phụ trách, đơn hàng, workflow, phạm vi xem."],
        ["6", "Jobs/đơn hàng", "Ngoài nước, trong nước, du học; quốc gia, nghề, số lượng, lương, phí, hạn, danh sách ứng viên."],
        ["7", "Workflow 20 bước", "Chuyển bước theo trách nhiệm, lịch sử, nhánh thi trượt/chọn lại Job, điều kiện hoàn thành."],
        ["8", "Đào tạo", "Hai track học tiếng và chuyên môn/nghề; tiến độ 0–100%; phiếu đánh giá tuần và minh chứng."],
        ["9", "Đại lý/CTV & hoa hồng", "Cây đại lý, CTV, cấu hình tỷ lệ, mốc hoa hồng, duyệt/chi, portal và thi đua doanh số."],
        ["10", "Tài chính", "Thu tiền 4 bước, khoản chi, phiếu thu/chi PDF, công nợ, phê duyệt, liên kết phát sinh hoa hồng."],
        ["11", "Vay & thu nợ", "Vay ngân hàng/nợ công ty, lịch trả góp, thu từng kỳ/thu hết, phiếu thu, gate hoàn tất."],
        ["12", "Visa & chuyến bay", "Hồ sơ visa, trạng thái/ngày phỏng vấn/kết quả, vé bay, điểm đến và nhắc xuất cảnh."],
        ["13", "Thông báo", "Nhắc việc 5 phút/lần, đa kênh In-app/Email/SMS/Zalo, điều hướng đến nghiệp vụ nguồn."],
        ["14", "Tin nhắn", "Danh bạ theo quan hệ, hội thoại, đính kèm/âm thanh, đọc, thu hồi và giới hạn người nhận."],
        ["15", "Trợ lý AI", "Hỏi đáp nghiệp vụ, phân tích hồ sơ, đọc ảnh/PDF, trích xuất CV; ngữ cảnh giới hạn theo user."],
        ["16", "Báo cáo & xuất file", "Dashboard, biểu đồ, phễu, tài chính, quốc gia, Excel/PDF/CSV và in phiếu."],
        ["17–20", "Hạ tầng kiểm soát", "Dashboard portal, MinIO hồ sơ, audit log, health check, Hangfire, backup/restore, deploy."],
    ]
    add_table(doc, ["#", "Phân hệ", "Năng lực nghiệp vụ"], module_rows, [1.2, 4.5, 11.2])

    add_heading(doc, "5. Các quy trình tài chính và hoa hồng", 1)
    add_heading(doc, "5.1 Thu tiền ứng viên", 2)
    add_table(doc, ["Đợt", "Tỷ lệ", "Mục đích", "Nguyên tắc kiểm soát"], [
        ["1. Đặt cọc", "20%", "Xác nhận tham gia", "Là mốc nghiệp vụ và có thể kích hoạt hoa hồng đặt cọc."],
        ["2. Phí dịch vụ", "30%", "Chi phí triển khai hồ sơ/dịch vụ", "Đóng theo thứ tự, không bỏ qua đợt trước."],
        ["3. Trước xuất cảnh", "30%", "Hoàn thiện nghĩa vụ trước bay", "Gắn với tiến độ gần xuất cảnh."],
        ["4. Tất toán", "20%", "Hoàn tất chi phí đơn hàng", "Đủ 100% hoặc có cam kết trả nợ hợp lệ để qua bước 15."],
    ], [3.3, 1.8, 5.2, 6.6])
    doc.add_paragraph("Mỗi khoản thu có mã, loại, số tiền, hạn thu và trạng thái Pending/Partial/Paid/Overdue/Refunded. Khi ghi nhận đã thu, dịch vụ chung xử lý thứ tự, actor, ngày, audit và kích hoạt hoa hồng để tránh lệch giữa các màn hình.")

    add_heading(doc, "5.2 Hoa hồng đại lý/CTV", 2)
    add_table(doc, ["Mốc", "Ý nghĩa", "Trạng thái xử lý"], [
        ["Đặt cọc", "Ghi nhận ứng viên đã cam kết", "Pending → Approved → Paid hoặc Cancelled"],
        ["Trúng tuyển", "Ứng viên đạt kết quả tuyển chọn", "Pending → Approved → Paid hoặc Cancelled"],
        ["Xuất cảnh", "Ứng viên hoàn thành mốc ra nước ngoài", "Pending → Approved → Paid hoặc Cancelled"],
    ], [3.3, 7.1, 6.5])
    add_bullets(doc, [
        "Hệ thống có khóa chống sinh trùng hoa hồng theo đại lý–ứng viên–mốc.",
        "Kế toán được quyền duyệt và chi theo quyết định nghiệp vụ hiện hành; đại lý/CTV có portal riêng để xem phần thuộc phạm vi.",
        "Tỷ lệ chia CTV hiện cần được snapshot tại thời điểm phát sinh để bảo toàn lịch sử — đây là hạng mục thay đổi đã được ghi nhận nhưng chưa hoàn tất.",
        "Chính sách hiển thị doanh số đối thủ trong bảng thi đua cần điều chỉnh: đại lý chỉ thấy thứ hạng của mình, các vai trò nội bộ vẫn xem đầy đủ."
    ])

    add_heading(doc, "5.3 Vay và thu nợ", 2)
    add_table(doc, ["Loại", "Cách quản lý", "Điều kiện hoàn tất"], [
        ["Vay ngân hàng", "Theo dõi hồ sơ vay, ngân hàng, số tiền, lãi suất, trạng thái giải ngân.", "Không dùng trạng thái tất toán nội bộ để chặn bước 20."],
        ["Nợ công ty", "Sinh lịch trả góp gốc+lãi, thu từng kỳ hoặc thu hết, mỗi lần thu sinh phiếu thu.", "Chỉ mở bước 20 khi đã thu đủ 100%; không miễn nợ, không cho tất toán thủ công khi còn dư."],
    ], [3.3, 7.5, 6.1])

    add_heading(doc, "6. Vai trò và mô hình kiểm soát", 1)
    roles = [
        ["Super Admin", "Toàn quyền; quản trị tài khoản/quyền; thao tác đặc biệt như đổi người/đổi đơn hàng."],
        ["Giám đốc", "Xem toàn hệ thống, báo cáo, audit; phê duyệt/giám sát hoa hồng và tài chính theo quyền."],
        ["Trưởng phòng tuyển dụng", "Quản Lead, ứng viên, cộng tác viên, vay, đào tạo; theo dõi dashboard/KPI."],
        ["NV tuyển dụng", "Tạo/chăm sóc Lead, tạo/sửa ứng viên, xem Job/đại lý, hỗ trợ vay."],
        ["Tư vấn viên", "Theo sát Lead và ứng viên phụ trách; đào tạo và hỗ trợ xuyên suốt."],
        ["Bộ phận hồ sơ", "Cập nhật hồ sơ, giấy tờ, tiến độ; xem Job, vay và visa."],
        ["Bộ phận Visa", "Quản visa, chuyến bay, cập nhật ứng viên liên quan."],
        ["Kế toán", "Toàn quyền nghiệp vụ thu, chi, phiếu, hoa hồng, vay/thu nợ; xem báo cáo."],
        ["Đại lý", "Chỉ xem ứng viên, CTV, hoa hồng và đào tạo trong đại lý; portal/thi đua riêng."],
        ["CTV", "Chỉ xem ứng viên và hoa hồng do mình giới thiệu; nhắn tin theo quan hệ."],
        ["Phụ huynh", "Chỉ xem hồ sơ, đào tạo, tài chính, vay của con/em; không thấy số liệu hoa hồng nhạy cảm."],
        ["Học viên", "Chỉ xem dữ liệu của chính mình; tài chính, vay, đào tạo, tin nhắn và AI cá nhân hóa."],
    ]
    add_table(doc, ["Vai trò", "Phạm vi chính"], roles, [4.1, 12.8])
    add_heading(doc, "6.1 Nguyên tắc kiểm soát", 2)
    add_bullets(doc, [
        "Kiểm soát hai lớp: quyền hành động (create/read/update/delete/approve) và phạm vi dữ liệu.",
        "Đại lý lọc theo AgentId; CTV theo CollaboratorId; phụ huynh/học viên theo đúng Candidate được liên kết.",
        "Trang chi tiết tiếp tục kiểm scope để chống truy cập trực tiếp bằng URL ngoài phạm vi.",
        "Thao tác nhạy cảm của super admin yêu cầu nhập lại mật khẩu và ghi audit.",
        "Tài khoản bị khóa được xoay security stamp để loại phiên cũ; production không seed tài khoản demo/mật khẩu mặc định."
    ])

    add_heading(doc, "7. Dashboard và bộ KPI dành cho Giám đốc", 1)
    kpis = [
        ["Phễu tuyển dụng", "Tổng Lead, Lead mới hôm nay, tỷ lệ chuyển đổi, ứng viên đang xử lý, tỷ lệ trúng tuyển, tỷ lệ đậu visa, tỷ lệ xuất cảnh", "Đo sức khỏe toàn chuỗi và xác định điểm rơi."],
        ["Năng lực cung ứng", "Job đang tuyển, ứng viên theo Job/quốc gia, sắp xuất cảnh 30 ngày", "Điều phối nguồn lực, kế hoạch xuất cảnh và đối tác."],
        ["Tài chính", "Doanh thu đã thu/tháng, tổng chi phí, lợi nhuận gộp, công nợ phải thu, khoản thu quá hạn", "Kiểm soát dòng tiền và chất lượng doanh thu."],
        ["Kênh & thị trường", "Lead theo nguồn, doanh thu theo quốc gia, tỉnh/thành và tỷ lệ chuyển đổi", "Phân bổ ngân sách marketing và ưu tiên thị trường."],
        ["Đại lý", "Top đại lý, doanh số, hoa hồng phải trả/đã trả", "Quản hiệu suất kênh và chi phí bán hàng."],
        ["Rủi ro vận hành", "Lead quá hạn chăm sóc, hồ sơ thiếu, visa/phỏng vấn đến hạn, khoản vay/chi/hoa hồng chờ xử lý", "Can thiệp sớm trước khi ảnh hưởng tiến độ và dòng tiền."],
    ]
    add_table(doc, ["Nhóm KPI", "Chỉ số hiện có", "Câu hỏi quản trị trả lời"], kpis, [3.5, 8.4, 5.0])
    add_heading(doc, "7.1 Nhịp điều hành đề xuất", 2)
    add_table(doc, ["Tần suất", "Nội dung họp/kiểm"], [
        ["Hàng ngày", "Lead mới, Lead quá hạn, lịch hẹn, hồ sơ/visa đến hạn, khoản thu quá hạn, chuyến bay 7–30 ngày."],
        ["Hàng tuần", "Phễu theo tư vấn viên/nguồn, tiến độ theo Job, đào tạo, công nợ, hoa hồng chờ duyệt/chi."],
        ["Hàng tháng", "Doanh thu–chi phí–lợi nhuận, hiệu quả quốc gia/đại lý/kênh, tỷ lệ chuyển đổi và năng suất nhân sự."],
        ["Hàng quý", "Chất lượng dữ liệu, phân quyền, audit, SLA, bảo mật, backup/restore và hiệu quả đầu tư hệ thống."],
    ], [3.1, 13.8])

    add_heading(doc, "8. Trạng thái QA và rủi ro còn lại", 1)
    qa_rows = [
        ["M01–M04", "Đăng nhập, quyền, tài khoản, CRM", "Đã sửa và verified code-level", "Còn cần runtime HTTP/UI đầy đủ."],
        ["M05–M08", "Ứng viên, Job, workflow, đào tạo", "Không có bug xác nhận; verified code-level", "Concurrency/lost update và E2E còn thiếu."],
        ["M09–M10", "Hoa hồng, tài chính", "Bug quan trọng đã sửa/verified", "Cần kiểm dữ liệu thật, phê duyệt chi và snapshot tỷ lệ CTV."],
        ["M11", "Vay & thu nợ", "Đã sửa; 88 test tổng đạt", "Migration mới chưa áp DB test; còn chờ xác minh độc lập/runtime."],
        ["M12", "Visa & Flight", "Bug attribution đã verified code-level", "Chưa có nút xác nhận đã bay; audit visa/flight chưa đủ."],
        ["M13", "Thông báo", "Sửa một phần; còn blocked requirement", "Chưa chốt CTV nhận gì và super admin/owner nhận thông báo tài chính thế nào."],
        ["M14–M20", "Chat, AI, báo cáo, dashboard, file, audit, deploy", "Có chức năng nhưng QA tổng thể pending", "Chưa đạt mức kết luận production-ready."],
    ]
    add_table(doc, ["Phạm vi", "Chức năng", "Tình trạng", "Khoảng trống chính"], qa_rows, [2.3, 4.0, 4.8, 5.8])
    add_heading(doc, "8.1 Rủi ro ưu tiên", 2)
    add_table(doc, ["Mức", "Rủi ro", "Tác động", "Khuyến nghị"], [
        ["Cao", "Public production khi còn tài khoản demo/mật khẩu mặc định hoặc secret chưa chuẩn", "Rò rỉ dữ liệu, chiếm quyền", "Chỉ deploy production bằng tài khoản super admin thật, TLS và secret môi trường."],
        ["Cao", "QA mới chủ yếu code-level; thiếu E2E đa vai trò và migration DB test", "Sai quyền, lỗi dữ liệu runtime", "UAT có kịch bản, DB test gần production, bằng chứng ảnh/log."],
        ["Cao", "Quy tắc người nhận thông báo M13 chưa chốt", "Sai người nhận hoặc lộ tiền hoa hồng", "Giám đốc phê duyệt ma trận recipient/nội dung."],
        ["Trung bình", "Nhiều entity chưa có concurrency token", "Hai người sửa đồng thời có thể ghi đè", "Bổ sung rowversion/optimistic concurrency cho hồ sơ, workflow, tài chính."],
        ["Trung bình", "Visa/flight thiếu audit và xác nhận ActualDepartureAt", "Báo cáo xuất cảnh thực tế có thể thiếu", "Thêm nút xác nhận đã bay + audit đầy đủ."],
        ["Trung bình", "Chat/AI/file/report chưa QA đầy đủ", "Rủi ro scope, file, dữ liệu AI, export", "Hoàn tất M14–M20 trước nghiệm thu production."],
    ], [1.8, 5.2, 4.6, 5.3])

    add_heading(doc, "9. Các quyết định Ban Giám đốc cần chốt", 1)
    decisions = [
        ["D1", "Thông báo tài chính", "Super admin có nhận cùng Accountant/Director không? Payment reminder ưu tiên owner ứng viên hay gửi thẳng phòng tài chính?", "M13"],
        ["D2", "Thông báo hoa hồng cho CTV", "Chỉ CTV trực tiếp hay toàn cây đại lý? Hiển thị tổng commission, phần chia CTV hay chỉ trạng thái không tiền?", "M13"],
        ["D3", "Luồng duyệt khoản chi", "Ai đề nghị, ai duyệt, ngưỡng tiền, có duyệt hai cấp hay không?", "M10"],
        ["D4", "Xác nhận xuất cảnh", "Ai được bấm 'Đã bay'; có cần bằng chứng và khóa sửa sau xác nhận?", "M12"],
        ["D5", "Hoa hồng CTV", "Chốt snapshot tỷ lệ tại thời điểm phát sinh và nguyên tắc điều chỉnh hồi tố.", "M09"],
        ["D6", "Bảng thi đua đại lý", "Xác nhận đại lý chỉ thấy thứ hạng của mình, nội bộ xem toàn bộ.", "M09"],
        ["D7", "Quyền xem đào tạo", "Xác nhận recruiter/document/visa/accountant đều được training:read.", "M08"],
        ["D8", "Tiêu chí go-live", "Phê duyệt checklist bắt buộc: UAT, bảo mật, backup/restore, migration, TLS, thông báo ngoài hệ thống, rollback.", "M20"],
    ]
    add_table(doc, ["Mã", "Chủ đề", "Nội dung cần quyết định", "Module"], decisions, [1.3, 4.0, 10.2, 1.5])

    add_heading(doc, "10. Lộ trình hoàn thiện đề xuất", 1)
    add_table(doc, ["Giai đoạn", "Thời gian", "Kết quả cần đạt"], [
        ["1. Khóa nghiệp vụ", "Tuần 1", "Chốt D1–D8; cập nhật đặc tả và ma trận quyền/recipient."],
        ["2. Hoàn thiện & migration", "Tuần 1–2", "M11 DB test; change request M08–M13; audit visa/flight; xác nhận đã bay; snapshot hoa hồng."],
        ["3. UAT xuyên suốt", "Tuần 2–3", "12 vai trò, 5 hồ sơ mẫu, các nhánh rớt thi/chuyển Job/vay/nợ/hoa hồng/visa/xuất cảnh."],
        ["4. QA M14–M20", "Tuần 3–4", "Chat, AI, báo cáo, file, audit, bảo mật và deployment có bằng chứng nghiệm thu."],
        ["5. Go-live kiểm soát", "Sau nghiệm thu", "Dữ liệu sạch, user thật, TLS, secret, backup, monitoring, đào tạo người dùng và phương án rollback."],
    ], [4.0, 2.6, 10.3])
    add_heading(doc, "10.1 Điều kiện go-live tối thiểu", 2)
    add_numbered(doc, [
        "Không còn tài khoản demo hoặc mật khẩu mặc định trên production; toàn bộ secret nằm trong biến môi trường.",
        "Migration chạy thành công trên bản sao DB test; có backup và thử restore thành công.",
        "UAT ký duyệt theo vai trò và theo quy trình end-to-end, gồm cả trường hợp lỗi và thao tác trái quyền.",
        "TLS/domain, health check, log, Hangfire, MinIO và cảnh báo vận hành hoạt động ổn định.",
        "Các quyết định M13 về thông báo và bảo mật nội dung hoa hồng đã được phê duyệt.",
        "Báo cáo tài chính, công nợ, xuất cảnh và hoa hồng đối soát đúng với dữ liệu mẫu đã biết trước."
    ])

    add_heading(doc, "11. Kiến trúc hỗ trợ vận hành (tóm tắt)", 1)
    add_table(doc, ["Lớp", "Công nghệ/năng lực", "Vai trò"], [
        ["Web", ".NET 10, Blazor Interactive Server, MudBlazor", "Giao diện responsive và xử lý nghiệp vụ server-side."],
        ["Xác thực", "ASP.NET Core Identity, Cookie, JWT, RBAC permission claims", "Đăng nhập, API và phân quyền hành động/phạm vi."],
        ["Dữ liệu", "PostgreSQL 16, EF Core 10", "Lưu dữ liệu nghiệp vụ và migration."],
        ["Hồ sơ", "MinIO/S3-compatible", "Lưu tài liệu, ảnh, đính kèm và URL tải có thời hạn."],
        ["Job nền", "Hangfire + PostgreSQL", "Quét nhắc việc định kỳ và điều phối gửi thông báo."],
        ["Báo cáo", "ClosedXML, QuestPDF, CSV", "Xuất Excel/PDF/CSV và in phiếu thu/chi."],
        ["AI", "Gemini", "Hỏi đáp, phân tích hồ sơ, đọc ảnh/PDF và trích xuất CV."],
        ["Triển khai", "Docker Compose, Caddy/Nginx, health check, backup/restore", "Vận hành local/production có reverse proxy và TLS."],
    ], [3.0, 7.0, 6.9])

    add_heading(doc, "12. Kết luận", 1)
    doc.add_paragraph(
        "POLYMIND OLMS đã phát triển từ một hệ thống lưu hồ sơ thành một nền tảng điều hành nghiệp vụ tương đối toàn diện. Các luồng CRM, ứng viên, đơn hàng, 20 bước tiến độ, đào tạo, tài chính, vay/nợ, hoa hồng, visa–xuất cảnh và cổng đối tác/cá nhân đã liên kết với nhau đủ để tạo giá trị quản trị rõ ràng."
    )
    doc.add_paragraph(
        "Tuy vậy, trạng thái hiện tại phù hợp nhất với giai đoạn hoàn thiện và UAT có kiểm soát, chưa nên tuyên bố production-ready cho toàn hệ thống. Trọng tâm tiếp theo không phải mở thêm nhiều chức năng, mà là khóa chính sách nghiệp vụ còn mở, xác minh runtime đa vai trò, kiểm thử dữ liệu/migration, hoàn thiện bảo mật–audit–backup và đối soát báo cáo."
    )
    add_callout(doc, "Thông điệp dành cho Giám đốc", "Hệ thống đã có đủ nền để chuẩn hóa vận hành và ra quyết định bằng dữ liệu. Nếu hoàn tất các quyết định D1–D8 và nghiệm thu theo lộ trình 30 ngày, POLYMIND có thể chuyển từ bản phát triển chức năng sang một nền tảng vận hành thực tế có kiểm soát.", LIGHT_ORANGE)

    add_heading(doc, "Phụ lục A — Các đường dẫn chức năng chính", 1)
    add_table(doc, ["Đường dẫn", "Chức năng"], [
        ["/", "Dashboard tổng quan nội bộ"], ["/leads", "Lead CRM"], ["/candidates", "Ứng viên/hồ sơ"],
        ["/jobs", "Jobs/đơn hàng"], ["/training", "Đào tạo"], ["/finance", "Tài chính & công nợ"],
        ["/loans", "Hỗ trợ vay"], ["/debt-collection", "Thu nợ công ty"], ["/agents", "Đại lý & hoa hồng/thi đua"],
        ["/visa", "Visa & xuất cảnh"], ["/reports", "Báo cáo"], ["/notifications", "Thông báo"],
        ["/messages", "Tin nhắn"], ["/ai", "Trợ lý AI"], ["/me", "Cổng học viên/phụ huynh"],
        ["/my-commissions", "Hoa hồng của đối tác"], ["/admin", "Quản trị, phân quyền, audit"], ["/guide", "Hướng dẫn sử dụng"],
    ], [5.0, 11.9])

    add_heading(doc, "Phụ lục B — Nguồn tham chiếu chính", 1)
    add_bullets(doc, [
        "Mã nguồn trong src/Polymind.Web, src/Polymind.Domain và src/Polymind.Infrastructure.",
        "docs/01-business-analysis.md; docs/03-workflow.md; docs/04-system-architecture.md.",
        "docs/testing/MODULE_QA_BOARD.md; docs/testing/SESSION_CHECKPOINT.md.",
        "docs/testing/modules/M01–M14: analysis, business flows, test cases, bug/fix/verification reports.",
        "WORKLOG.md đến phiên 68 và kết quả dotnet test ngày 11/07/2026."
    ])
    doc.add_paragraph("— Hết báo cáo —").alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


if __name__ == "__main__":
    OUT.parent.mkdir(parents=True, exist_ok=True)
    report = build_report()
    report.save(OUT)
    print(OUT)
